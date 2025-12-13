from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Body
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import Base, get_db
from app.core.file_service import FileService
from app.core.query import CursorPaginationResult, QueryInput
from app.domain.services.file_service import FileQueryService
from app.presentation.api.dependencies import get_current_user
from app.presentation.api.v1.schemas.auth import TokenData
from app.utils.table_lookup import get_table_with_schema
from typing import List, Optional, Dict, Any
from pydantic import BaseModel, Field
from datetime import datetime, timedelta
from dateutil import parser as date_parser
import os
import docx
import openpyxl
import csv
import codecs

router = APIRouter()

def parse_datetime_safe(datetime_str: str) -> datetime:
    """
    Parse datetime string and convert to timezone-naive datetime.
    
    Args:
        datetime_str: Datetime string to parse
        
    Returns:
        timezone-naive datetime object
        
    Raises:
        ValueError: If datetime string is invalid
    """
    try:
        parsed_time = date_parser.parse(datetime_str)
        # Convert to timezone-naive datetime if it has timezone info
        if parsed_time.tzinfo is not None:
            return parsed_time.replace(tzinfo=None)
        else:
            return parsed_time
    except Exception as e:
        raise ValueError(f"Invalid datetime format: {str(e)}")

# Request/Response schemas for extract-file-content endpoint
class ExtractFileContentRequest(BaseModel):
    file_path: str = Field(..., description="Đường dẫn đến file cần trích xuất nội dung")

class ExtractFileContentResponse(BaseModel):
    file_path: str = Field(..., description="Đường dẫn file")
    file_name: str = Field(..., description="Tên file")
    file_size: int = Field(..., description="Kích thước file (bytes)")
    file_extension: str = Field(..., description="Phần mở rộng file")
    modified_time: str = Field(..., description="Thời gian chỉnh sửa cuối")
    content: str = Field(..., description="Nội dung được trích xuất")
    content_length: int = Field(..., description="Độ dài nội dung")
    success: bool = Field(default=True, description="Trạng thái thành công")
    message: str = Field(default="Trích xuất nội dung file thành công", description="Thông báo")

# Dashboard response schema
class FileDashboardResponse(BaseModel):
    total_files: int = Field(..., description="Tổng số lượng file")
    total_size: int = Field(..., description="Tổng dung lượng file (bytes)")
    total_size_mb: float = Field(..., description="Tổng dung lượng file (MB)")
    processed_files: int = Field(..., description="Số lượng file đã xử lý")
    unprocessed_files: int = Field(..., description="Số lượng file chưa xử lý")
    processing_rate: float = Field(..., description="Tỷ lệ xử lý (%)")
    avg_processing_duration: Optional[float] = Field(None, description="Thời gian xử lý trung bình (giây)")
    files_by_extension: dict = Field(..., description="Số lượng file theo phần mở rộng")
    files_by_status: dict = Field(..., description="Số lượng file theo trạng thái xử lý")

# Period statistics request schema
class PeriodStatsRequest(BaseModel):
    period: str = Field(..., description="Kỳ thống kê: 'day', 'month', 'quarter', 'year'")
    from_time: Optional[str] = Field(None, description="Thời gian bắt đầu (YYYY-MM-DD hoặc YYYY-MM-DD HH:MM:SS)")
    to_time: Optional[str] = Field(None, description="Thời gian kết thúc (YYYY-MM-DD hoặc YYYY-MM-DD HH:MM:SS)")

# Period statistics response schema
class PeriodStatsResponse(BaseModel):
    period: str = Field(..., description="Kỳ thống kê")
    from_time: Optional[str] = Field(None, description="Thời gian bắt đầu")
    to_time: Optional[str] = Field(None, description="Thời gian kết thúc")
    total_files: int = Field(..., description="Tổng số lượng file")
    total_size: int = Field(..., description="Tổng dung lượng file (bytes)")
    total_size_mb: float = Field(..., description="Tổng dung lượng file (MB)")
    statistics: List[Dict[str, Any]] = Field(..., description="Thống kê chi tiết theo kỳ")
    files_by_extension: dict = Field(..., description="Số lượng file theo phần mở rộng")
    files_by_status: dict = Field(..., description="Số lượng file theo trạng thái xử lý")

# Country and Technology statistics request schema
class CountryTechStatsRequest(BaseModel):
    from_time: Optional[str] = Field(None, description="Thời gian bắt đầu (YYYY-MM-DD hoặc YYYY-MM-DD HH:MM:SS)")
    to_time: Optional[str] = Field(None, description="Thời gian kết thúc (YYYY-MM-DD hoặc YYYY-MM-DD HH:MM:SS)")
    sort_by: str = Field(default="count", description="Sắp xếp theo: 'count' hoặc 'name'")
    sort_order: str = Field(default="desc", description="Thứ tự sắp xếp: 'asc' hoặc 'desc'")
    limit: Optional[int] = Field(None, description="Giới hạn số lượng kết quả trả về")

# Country and Technology statistics response schema
class CountryTechStatsResponse(BaseModel):
    from_time: Optional[str] = Field(None, description="Thời gian bắt đầu")
    to_time: Optional[str] = Field(None, description="Thời gian kết thúc")
    total_files: int = Field(..., description="Tổng số lượng file trong khoảng thời gian")
    listed_nations: List[Dict[str, Any]] = Field(..., description="Danh sách quốc gia và số lượng tài liệu")
    listed_technologies: List[Dict[str, Any]] = Field(..., description="Danh sách công nghệ và số lượng tài liệu")
    total_nations: int = Field(..., description="Tổng số quốc gia unique")
    total_technologies: int = Field(..., description="Tổng số công nghệ unique")

@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    fields: Optional[List[str]] = Query(None, description="Fields to include in response"),
    is_save: bool = Query(True, description="If False, only save file to disk without saving to database"),
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Upload a file.
    
    The file will be:
    1. Hashed to detect duplicates
    2. Saved to static/uploads directory
    3. Metadata stored in database (if is_save=True)
    4. Returns file information
    
    Parameters:
    - file: The file to upload
    - fields: Optional list of fields to include in response. Use "*" for all fields.
    - is_save: If True (default), save to database. If False, only save to disk.
    """
    try:
        file_service = FileService(session)
        result = await file_service.save_file(file, current_user.user_id, fields, is_save=is_save)
        return result
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error uploading file: {str(e)}"
        ) 
        

@router.post("/my-files", response_model=CursorPaginationResult)
async def query_with_cursor(
    query_input: QueryInput,
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Query files with cursor-based pagination.
    If include_children is True, will include files from current user and all child users.
    
    Parameters:
    - query_input: Query parameters including:
        - table_name: Name of the model/table to query
        - ids: List of IDs to filter by
        - fields: List of fields to select
        - page: Page number (1-based)
        - page_size: Number of items per page
        - cursor: Cursor for pagination
        - sort_by: Field to sort by
        - sort_order: Sort order ("asc" or "desc")
        - condition: Dictionary of conditions to filter by
        - include_children: If True, include files from current user and all child users
    """
    try:
        # Get model class from SQLAlchemy metadata with schema support
        model = get_table_with_schema(query_input.table_name)
        
        # Create FileQueryService instance with session
        file_query_service = FileQueryService(session)
        
        # Execute query with user hierarchy support
        result = await file_query_service.query_with_cursor(model, query_input, current_user.user_id)
        return result
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error executing query: {str(e)}"
        )

@router.post("/my-with-children", response_model=CursorPaginationResult)
async def query_files_with_children(
    query_input: QueryInput,
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Query files from current user and all their child users based on role hierarchy.
    
    Parameters:
    - query_input: Query parameters including:
        - table_name: Name of the model/table to query
        - ids: List of IDs to filter by
        - fields: List of fields to select
        - page: Page number (1-based)
        - page_size: Number of items per page
        - cursor: Cursor for pagination
        - sort_by: Field to sort by
        - sort_order: Sort order ("asc" or "desc")
        - condition: Dictionary of conditions to filter by
        - search_text: Text to search for
        - search_fields: List of fields to search in
    """
    try:
        # Get model class from SQLAlchemy metadata with schema support
        model = get_table_with_schema(query_input.table_name)
        
        # Create FileQueryService instance with session
        file_query_service = FileQueryService(session)
        
        # Execute query specifically for files with children
        result = await file_query_service.query_user_files_with_children(model, query_input, current_user.user_id)
        return result
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error executing query: {str(e)}"
        )


# File content extraction functions
def extract_docx_content(file_path: str) -> str:
    """Trích xuất nội dung từ file DOCX"""
    try:
        doc = docx.Document(file_path)
        content = []
        
        # Trích xuất text từ các paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Trích xuất text từ các bảng
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file DOCX: {str(e)}")


def extract_doc_content(file_path: str) -> str:
    """Trích xuất nội dung từ file DOC (cần python-docx2txt hoặc antiword)"""
    try:
        # Thử sử dụng python-docx2txt nếu có
        try:
            import docx2txt
            content = docx2txt.process(file_path)
            return content if content else "Không thể trích xuất nội dung từ file DOC"
        except ImportError:
            return "Cần cài đặt thư viện docx2txt để đọc file DOC"
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file DOC: {str(e)}")


def extract_xlsx_content(file_path: str) -> str:
    """Trích xuất nội dung từ file XLSX"""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        content = []
        
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            content.append(f"=== Sheet: {sheet_name} ===")
            
            for row in worksheet.iter_rows(values_only=True):
                # Lọc các cell không rỗng
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data if cell):
                    content.append(" | ".join(row_data))
        
        workbook.close()
        return "\n".join(content)
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file XLSX: {str(e)}")


def extract_text_content(file_path: str) -> str:
    """Trích xuất nội dung từ file TXT hoặc DAT"""
    try:
        # Thử các encoding khác nhau
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # Nếu tất cả encoding đều thất bại, thử đọc dưới dạng binary
        with open(file_path, 'rb') as file:
            raw_content = file.read()
            # Thử decode với errors='ignore'
            content = raw_content.decode('utf-8', errors='ignore')
            return content
            
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file text: {str(e)}")


def extract_csv_content(file_path: str) -> str:
    """Trích xuất nội dung từ file CSV"""
    try:
        content = []
        
        # Thử các encoding khác nhau
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as csvfile:
                    # Tự động phát hiện delimiter
                    sample = csvfile.read(1024)
                    csvfile.seek(0)
                    
                    delimiter = ','
                    if '\t' in sample:
                        delimiter = '\t'
                    elif ';' in sample:
                        delimiter = ';'
                    
                    csv_reader = csv.reader(csvfile, delimiter=delimiter)
                    
                    for row_num, row in enumerate(csv_reader, 1):
                        if row and any(cell.strip() for cell in row):
                            content.append(f"Row {row_num}: {' | '.join(row)}")
                    
                    return "\n".join(content)
                    
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        raise Exception("Không thể đọc file CSV với bất kỳ encoding nào")
        
    except Exception as e:
        raise Exception(f"Lỗi khi đọc file CSV: {str(e)}")


@router.post("/extract-file-content", response_model=ExtractFileContentResponse)
async def extract_file_content(
    request: ExtractFileContentRequest,
    current_user: TokenData = Depends(get_current_user)
):
    """
    API để trích xuất nội dung từ file trong thư mục static
    Hỗ trợ các định dạng: doc, docx, xlsx, txt, csv, dat
    """
    try:
        file_path = request.file_path
        
        # Kiểm tra file_path có chứa đường dẫn đầy đủ tới thư mục static
        if not file_path.startswith('static'):
            file_path = f"static/{file_path}"
        
        # Kiểm tra file có tồn tại không
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=404,
                detail=f"File {file_path} không tồn tại"
            )
        
        # Lấy phần mở rộng của file
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Danh sách các phần mở rộng được hỗ trợ
        supported_extensions = ['.doc', '.docx', '.xlsx', '.txt', '.csv', '.dat']
        
        if file_extension not in supported_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Định dạng file {file_extension} không được hỗ trợ. Chỉ hỗ trợ: {', '.join(supported_extensions)}"
            )
        
        content = ""
        
        # Trích xuất nội dung dựa trên định dạng file
        if file_extension == '.docx':
            content = extract_docx_content(file_path)
        elif file_extension == '.doc':
            content = extract_doc_content(file_path)
        elif file_extension == '.xlsx':
            content = extract_xlsx_content(file_path)
        elif file_extension in ['.txt', '.dat']:
            content = extract_text_content(file_path)
        elif file_extension == '.csv':
            content = extract_csv_content(file_path)
        
        # Lấy thông tin file
        file_stats = os.stat(file_path)
        file_info = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size": file_stats.st_size,
            "file_extension": file_extension,
            "modified_time": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "content": content,
            "content_length": len(content)
        }
        
        return file_info
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Lỗi khi trích xuất nội dung file: {str(e)}"
        )


@router.get("/dashboard", response_model=FileDashboardResponse)
async def get_file_dashboard(
    from_time: Optional[datetime] = Query(None, description="Thời gian bắt đầu (ISO format)"),
    to_time: Optional[datetime] = Query(None, description="Thời gian kết thúc (ISO format)"),
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    API để lấy thống kê dashboard cho file
    Bao gồm: tổng số file, dung lượng, số file đã xử lý/chưa xử lý, thống kê theo extension
    Có thể lọc theo khoảng thời gian với from_time và to_time
    """
    try:
        from sqlalchemy import func, case, select
        from app.domain.models import File
        
        # Lấy danh sách user IDs trong hierarchy
        file_query_service = FileQueryService(session)
        user_ids = await file_query_service.get_user_hierarchy_ids(current_user.user_id)
        
        # Xây dựng điều kiện lọc theo thời gian
        time_conditions = []
        if from_time:
            time_conditions.append(File.created_at >= from_time)
        if to_time:
            time_conditions.append(File.created_at <= to_time)
        
        # Tổng số file và tổng dung lượng
        total_query_conditions = [File.created_by.in_(user_ids)] + time_conditions
        total_query = await session.execute(
            select(
                func.count(File.id).label('total_files'),
                func.coalesce(func.sum(File.size), 0).label('total_size')
            ).where(*total_query_conditions)
        )
        total_result = total_query.first()
        total_files = total_result.total_files or 0
        total_size = total_result.total_size or 0
        
        # Số file đã xử lý và chưa xử lý
        processed_query_conditions = [File.created_by.in_(user_ids)] + time_conditions
        processed_query = await session.execute(
            select(
                func.count(case((File.is_processed == True, 1))).label('processed'),
                func.count(case((File.is_processed == False, 1))).label('unprocessed'),
                func.count(case((File.is_processed.is_(None), 1))).label('pending')
            ).where(*processed_query_conditions)
        )
        processed_result = processed_query.first()
        processed_files = processed_result.processed or 0
        unprocessed_files = (processed_result.unprocessed or 0) + (processed_result.pending or 0)
        
        # Thống kê theo extension
        extension_query_conditions = [File.created_by.in_(user_ids)] + time_conditions
        extension_query = await session.execute(
            select(
                File.extension,
                func.count(File.id).label('count')
            ).where(*extension_query_conditions)
            .group_by(File.extension)
            .order_by(func.count(File.id).desc())
        )
        extension_results = extension_query.all()
        files_by_extension = {row.extension or 'unknown': row.count for row in extension_results}
        
        # Thống kê theo trạng thái xử lý
        status_query_conditions = [File.created_by.in_(user_ids)] + time_conditions
        status_query = await session.execute(
            select(
                case(
                    (File.is_processed == True, 'processed'),
                    (File.is_processed == False, 'failed'),
                    (File.is_processed.is_(None), 'pending')
                ).label('status'),
                func.count(File.id).label('count')
            ).where(*status_query_conditions)
            .group_by(File.is_processed)
        )
        status_results = status_query.all()
        files_by_status = {row.status: row.count for row in status_results}
        
        # Tính thời gian xử lý trung bình (chỉ cho các file đã xử lý có processing_duration)
        avg_duration_query_conditions = [
            File.created_by.in_(user_ids),
            File.is_processed == True,
            File.processing_duration.isnot(None)
        ] + time_conditions
        avg_duration_query = await session.execute(
            select(
                func.avg(File.processing_duration).label('avg_duration')
            ).where(*avg_duration_query_conditions)
        )
        avg_duration_result = avg_duration_query.first()
        avg_processing_duration = round(avg_duration_result.avg_duration, 2) if avg_duration_result.avg_duration else None
        
        # Tính tỷ lệ xử lý
        processing_rate = (processed_files / total_files * 100) if total_files > 0 else 0
        
        # Chuyển đổi dung lượng sang MB
        total_size_mb = round(total_size / (1024 * 1024), 2)
        
        return FileDashboardResponse(
            total_files=total_files,
            total_size=total_size,
            total_size_mb=total_size_mb,
            processed_files=processed_files,
            unprocessed_files=unprocessed_files,
            processing_rate=round(processing_rate, 2),
            avg_processing_duration=avg_processing_duration,
            files_by_extension=files_by_extension,
            files_by_status=files_by_status
        )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Lỗi khi lấy thống kê dashboard: {str(e)}"
        )


@router.post("/period-stats", response_model=PeriodStatsResponse)
async def get_period_statistics(
    request: PeriodStatsRequest,
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    API để lấy thống kê file theo kỳ (ngày, tháng, quý, năm)
    
    Parameters:
    - period: Kỳ thống kê ('day', 'month', 'quarter', 'year')
    - from_time: Thời gian bắt đầu (optional)
    - to_time: Thời gian kết thúc (optional)
    """
    try:
        from sqlalchemy import func, case, select, extract, text
        from app.domain.models import File
        
        # Validate period parameter
        valid_periods = ['day', 'month', 'quarter', 'year']
        if request.period not in valid_periods:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid period. Must be one of: {', '.join(valid_periods)}"
            )
        
        # Parse time parameters
        from_time = None
        to_time = None
        
        if request.from_time:
            try:
                from_time = parse_datetime_safe(request.from_time)
                print(f"Parsed from_time: {from_time} (timezone-naive: {from_time.tzinfo is None})")
            except ValueError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid from_time format: {str(e)}"
                )
        
        if request.to_time:
            try:
                to_time = parse_datetime_safe(request.to_time)
                print(f"Parsed to_time: {to_time} (timezone-naive: {to_time.tzinfo is None})")
            except ValueError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid to_time format: {str(e)}"
                )
        
        # Lấy danh sách user IDs trong hierarchy
        file_query_service = FileQueryService(session)
        user_ids = await file_query_service.get_user_hierarchy_ids(current_user.user_id)
        
        # Base query with user hierarchy filter
        base_query = select(File).where(File.created_by.in_(user_ids))
        
        # Add time filters if provided
        if from_time:
            base_query = base_query.where(File.created_at >= from_time)
        if to_time:
            base_query = base_query.where(File.created_at <= to_time)
        
        # Build period grouping based on period type
        if request.period == 'day':
            period_expr = func.date(File.created_at)
            period_label = 'date'
        elif request.period == 'month':
            period_expr = func.date_trunc('month', File.created_at)
            period_label = 'month'
        elif request.period == 'quarter':
            period_expr = func.date_trunc('quarter', File.created_at)
            period_label = 'quarter'
        elif request.period == 'year':
            period_expr = func.date_trunc('year', File.created_at)
            period_label = 'year'
        
        # Get period statistics
        stats_query = await session.execute(
            select(
                period_expr.label('period_date'),
                func.count(File.id).label('file_count'),
                func.coalesce(func.sum(File.size), 0).label('total_size')
            )
            .where(File.created_by.in_(user_ids))
            .where(from_time <= File.created_at if from_time else True)
            .where(File.created_at <= to_time if to_time else True)
            .group_by(period_expr)
            .order_by(period_expr)
        )
        
        stats_results = stats_query.all()
        
        # Format statistics data
        statistics = []
        total_files = 0
        total_size = 0
        
        for row in stats_results:
            period_date = row.period_date
            file_count = row.file_count
            size = row.total_size
            
            total_files += file_count
            total_size += size
            
            # Format period display based on type
            if request.period == 'day':
                period_display = period_date.strftime('%Y-%m-%d')
            elif request.period == 'month':
                period_display = period_date.strftime('%Y-%m')
            elif request.period == 'quarter':
                year = period_date.year
                quarter = (period_date.month - 1) // 3 + 1
                period_display = f"{year}-Q{quarter}"
            elif request.period == 'year':
                period_display = str(period_date.year)
            
            statistics.append({
                period_label: period_display,
                'file_count': file_count,
                'total_size': size,
                'total_size_mb': round(size / (1024 * 1024), 2)
            })
        
        # Get files by extension for the period
        extension_query = await session.execute(
            select(
                File.extension,
                func.count(File.id).label('count')
            )
            .where(File.created_by.in_(user_ids))
            .where(from_time <= File.created_at if from_time else True)
            .where(File.created_at <= to_time if to_time else True)
            .group_by(File.extension)
            .order_by(func.count(File.id).desc())
        )
        extension_results = extension_query.all()
        files_by_extension = {row.extension or 'unknown': row.count for row in extension_results}
        
        # Get files by status for the period
        status_query = await session.execute(
            select(
                case(
                    (File.is_processed == True, 'processed'),
                    (File.is_processed == False, 'failed'),
                    (File.is_processed.is_(None), 'pending')
                ).label('status'),
                func.count(File.id).label('count')
            )
            .where(File.created_by.in_(user_ids))
            .where(from_time <= File.created_at if from_time else True)
            .where(File.created_at <= to_time if to_time else True)
            .group_by(File.is_processed)
        )
        status_results = status_query.all()
        files_by_status = {row.status: row.count for row in status_results}
        
        # Calculate total size in MB
        total_size_mb = round(total_size / (1024 * 1024), 2)
        
        return PeriodStatsResponse(
            period=request.period,
            from_time=request.from_time,
            to_time=request.to_time,
            total_files=total_files,
            total_size=total_size,
            total_size_mb=total_size_mb,
            statistics=statistics,
            files_by_extension=files_by_extension,
            files_by_status=files_by_status
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Lỗi khi lấy thống kê theo kỳ: {str(e)}"
        )


@router.post("/country-tech-stats", response_model=CountryTechStatsResponse)
async def get_country_technology_statistics(
    request: CountryTechStatsRequest,
    current_user: TokenData = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    API để lấy thống kê tần suất số lượng tài liệu nhắc đến các quốc gia và công nghệ
    trong khoảng thời gian từ from_time đến to_time
    
    Parameters:
    - from_time: Thời gian bắt đầu (optional)
    - to_time: Thời gian kết thúc (optional)
    - sort_by: Sắp xếp theo 'count' hoặc 'name' (default: 'count')
    - sort_order: Thứ tự sắp xếp 'asc' hoặc 'desc' (default: 'desc')
    - limit: Giới hạn số lượng kết quả trả về (optional)
    """
    try:
        from sqlalchemy import func, select, text
        from app.domain.models import File
        
        # Parse time parameters
        from_time = None
        to_time = None
        
        if request.from_time:
            try:
                from_time = parse_datetime_safe(request.from_time)
            except ValueError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid from_time format: {str(e)}"
                )
        
        if request.to_time:
            try:
                to_time = parse_datetime_safe(request.to_time)
            except ValueError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid to_time format: {str(e)}"
                )
        
        # Validate sort parameters
        if request.sort_by not in ['count', 'name']:
            raise HTTPException(
                status_code=400,
                detail="sort_by must be 'count' or 'name'"
            )
        
        if request.sort_order not in ['asc', 'desc']:
            raise HTTPException(
                status_code=400,
                detail="sort_order must be 'asc' or 'desc'"
            )
        
        # Lấy danh sách user IDs trong hierarchy
        file_query_service = FileQueryService(session)
        user_ids = await file_query_service.get_user_hierarchy_ids(current_user.user_id)
        
        # Base query with user hierarchy filter
        base_query = select(File).where(File.created_by.in_(user_ids))
        
        # Add time filters if provided
        if from_time:
            base_query = base_query.where(File.created_at >= from_time)
        if to_time:
            base_query = base_query.where(File.created_at <= to_time)
        
        # Get total files count
        total_files_query = select(func.count(File.id)).where(File.created_by.in_(user_ids))
        if from_time:
            total_files_query = total_files_query.where(File.created_at >= from_time)
        if to_time:
            total_files_query = total_files_query.where(File.created_at <= to_time)
        
        total_files_result = await session.execute(total_files_query)
        total_files = total_files_result.scalar() or 0
        
        # Get listed_nations statistics
        nations_query = select(
            func.unnest(File.listed_nation).label('nation'),
            func.count(File.id).label('count')
        ).where(File.created_by.in_(user_ids)) \
         .where(File.listed_nation.isnot(None)) \
         .where(func.array_length(File.listed_nation, 1) > 0)
        
        if from_time:
            nations_query = nations_query.where(File.created_at >= from_time)
        if to_time:
            nations_query = nations_query.where(File.created_at <= to_time)
            
        nations_query = nations_query.group_by(func.unnest(File.listed_nation)) \
         .order_by(
             func.count(File.id).desc() if request.sort_by == 'count' and request.sort_order == 'desc'
             else func.count(File.id).asc() if request.sort_by == 'count' and request.sort_order == 'asc'
             else func.unnest(File.listed_nation).asc() if request.sort_by == 'name' and request.sort_order == 'asc'
             else func.unnest(File.listed_nation).desc()
         )
        
        # Apply limit if specified
        if request.limit:
            nations_query = nations_query.limit(request.limit)
        
        nations_results = (await session.execute(nations_query)).all()
        listed_nations = [
            {
                "name": row.nation,
                "count": row.count,
                "percentage": round((row.count / total_files * 100), 2) if total_files > 0 else 0
            }
            for row in nations_results
        ]
        
        # Get listed_technologies statistics
        technologies_query = select(
            func.unnest(File.listed_technology).label('technology'),
            func.count(File.id).label('count')
        ).where(File.created_by.in_(user_ids)) \
         .where(File.listed_technology.isnot(None)) \
         .where(func.array_length(File.listed_technology, 1) > 0)
        
        if from_time:
            technologies_query = technologies_query.where(File.created_at >= from_time)
        if to_time:
            technologies_query = technologies_query.where(File.created_at <= to_time)
            
        technologies_query = technologies_query.group_by(func.unnest(File.listed_technology)) \
         .order_by(
             func.count(File.id).desc() if request.sort_by == 'count' and request.sort_order == 'desc'
             else func.count(File.id).asc() if request.sort_by == 'count' and request.sort_order == 'asc'
             else func.unnest(File.listed_technology).asc() if request.sort_by == 'name' and request.sort_order == 'asc'
             else func.unnest(File.listed_technology).desc()
         )
        
        # Apply limit if specified
        if request.limit:
            technologies_query = technologies_query.limit(request.limit)
        
        technologies_results = (await session.execute(technologies_query)).all()
        listed_technologies = [
            {
                "name": row.technology,
                "count": row.count,
                "percentage": round((row.count / total_files * 100), 2) if total_files > 0 else 0
            }
            for row in technologies_results
        ]
        
        # Get unique counts
        total_nations = len(listed_nations)
        total_technologies = len(listed_technologies)
        
        return CountryTechStatsResponse(
            from_time=request.from_time,
            to_time=request.to_time,
            total_files=total_files,
            listed_nations=listed_nations,
            listed_technologies=listed_technologies,
            total_nations=total_nations,
            total_technologies=total_technologies
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Lỗi khi lấy thống kê quốc gia và công nghệ: {str(e)}"
        )


