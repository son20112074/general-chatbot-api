from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
import os
import re
from bs4 import BeautifulSoup
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.domain.models.task import Task
from app.domain.models.task_work import TaskWork
from app.domain.models.user import User


class WorkIssuesDocxCreator:
    def __init__(self, session):
        self.session = session
    
    async def create_docx(self, work_list, download_options=[]):
        self.document = Document()
        self.write_title(
            "DANH SÁCH VẤN ĐỀ TRONG QUÁ TRÌNH THỰC HIỆN CÔNG VIỆC", 
            "Thông tin chi tiết của {} vấn đề trong quá trình thực hiện công việc".format(len(work_list))
        ) 
        self.write_paragraph("", '')
        
        # Process each work
        for i, work in enumerate(work_list, 1):
            await self.process_work(work, str(i))
        
        file_name = '{}_{}.docx'.format('danh_sach_van_de_cong_viec', str(datetime.today().timestamp()))
        
        # Create save path with year/month/day structure
        current_date = datetime.now()
        year = str(current_date.year)
        month = str(current_date.month).zfill(2)  # Zero-padded month
        day = str(current_date.day).zfill(2)      # Zero-padded day
        
        save_path = f'static/downloads/docx/{year}/{month}/{day}'
        
        # Create directory structure if it doesn't exist
        os.makedirs(save_path, exist_ok=True)

        self.document.save(save_path + '/' + file_name)
        
        return settings.BASE_URL + '/' + save_path + '/' + file_name

    async def process_work(self, work, index=""):
        """Process a single work with all its information"""
        print(f"Processing work {index}: {work.get('title', 'No title')} (ID: {work.get('id', 'No ID')})")
        
        # Write work title with index
        work_title = f"{index}. {work.get('title', 'Công việc không có tiêu đề')}"
        self.write_block(work_title, '', [])
        
        # Write work description
        work_description = work.get('description', 'Không có mô tả')
        self.write_paragraph('Mô tả: ', work_description)
        
        # Write work content (convert HTML to text)
        work_content = work.get('content', 'Không có nội dung')
        content_text = self.html_to_text(work_content) if work_content else "Không có nội dung"
        self.write_paragraph('Nội dung: ', content_text)
        
        # Write creator information
        creator = work.get('creator')
        if creator:
            creator_name = creator.get('full_name', 'Không xác định')
            creator_email = creator.get('email', 'Không có email')
            self.write_paragraph('Người thực hiện: ', f"{creator_name} ({creator_email})")
        else:
            self.write_paragraph('Người thực hiện: ', 'Không xác định')
        
        # Write creation date
        created_at = work.get('created_at')
        if created_at:
            # Convert ISO string to datetime if needed
            if isinstance(created_at, str):
                try:
                    created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                except:
                    pass
            
            if isinstance(created_at, datetime):
                created_date_str = created_at.strftime('%d/%m/%Y %H:%M')
            else:
                created_date_str = str(created_at)
            self.write_paragraph('Ngày thực hiện: ', created_date_str)
        else:
            self.write_paragraph('Ngày thực hiện: ', 'Không xác định')
        
        # Write task information
        task = work.get('task')
        if task:
            task_id = task.get('id', 'Không có ID')
            task_title = task.get('title', 'Không có tiêu đề')
            self.write_paragraph('Nhiệm vụ: ', f"(Mã {task_id}) - {task_title}")
            
            # Write task description
            task_description = task.get('description', 'Không có mô tả')
            self.write_paragraph('Mô tả nhiệm vụ: ', task_description)
            
            # Write task status
            task_status = task.get('status', 'new')
            status_vn = self.get_status_vietnamese(task_status)
            self.write_paragraph('Trạng thái nhiệm vụ: ', status_vn)
            
            # Write task dates
            task_start_date = task.get('start_date')
            if task_start_date:
                if isinstance(task_start_date, str):
                    try:
                        task_start_date = datetime.fromisoformat(task_start_date.replace('Z', '+00:00'))
                    except:
                        pass
                
                if isinstance(task_start_date, datetime):
                    start_date_str = task_start_date.strftime('%d/%m/%Y')
                else:
                    start_date_str = str(task_start_date)
                self.write_paragraph('Thời gian bắt đầu nhiệm vụ: ', start_date_str)
            
            task_end_date = task.get('end_date')
            if task_end_date:
                if isinstance(task_end_date, str):
                    try:
                        task_end_date = datetime.fromisoformat(task_end_date.replace('Z', '+00:00'))
                    except:
                        pass
                        
                if isinstance(task_end_date, datetime):
                    end_date_str = task_end_date.strftime('%d/%m/%Y')
                else:
                    end_date_str = str(task_end_date)
                self.write_paragraph('Thời gian kết thúc nhiệm vụ: ', end_date_str)
        else:
            self.write_paragraph('Nhiệm vụ: ', 'Không có thông tin nhiệm vụ')
        
        # Write other attributes
        attributes = work.get('attributes', {})
        if attributes:
            self.write_paragraph('Thông tin khác: ', '')
            for key, value in attributes.items():
                self.write_paragraph(f'  {key}: ', str(value))
        
        # Write image links if any
        # image_links = work.get('image_links', [])
        # if image_links:
        #     self.write_paragraph('Liên kết hình ảnh: ', '')
        #     for i, link in enumerate(image_links, 1):
        #         self.write_paragraph(f'  {i}. ', link)
        
        # # Write file links if any
        # file_links = work.get('file_links', [])
        # if file_links:
        #     self.write_paragraph('Liên kết tệp: ', '')
        #     for i, link in enumerate(file_links, 1):
        #         self.write_paragraph(f'  {i}. ', link)
        
        # Write difficulty flag
        # is_difficult = work.get('is_difficult', False)
        # difficulty_text = "Có" if is_difficult else "Không"
        # self.write_paragraph('Công việc khó: ', difficulty_text)
        
        # Add separator between works
        self.write_paragraph("", "")

    def html_to_text(self, html_content):
        """Convert HTML content to plain text"""
        if not html_content:
            return ""
            
        try:
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            return text
            
        except Exception as e:
            print(f"Error converting HTML to text: {e}")
            # Fallback: return original content with basic HTML tag removal
            return re.sub(r'<[^>]+>', '', html_content)

    def set_empty_to_none(self, info={}):
        for key, value in info.items():
            if value == None:
                info[key] = ''
        # pass


    def write_title(self, title, summary):
        self.document.add_heading('', 0).add_run(title).font.size = Pt(16)
        p = self.document.add_paragraph('').add_run(summary).italic = True


    def write_heading_bullet(self, title):
        self.document.add_paragraph(
            title, style='List Number'
        )
        
    def write_paragraph(self, title_bold, content):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(title_bold).bold = True
        paragraph.add_run(content).italic = True
        # paragraph.alignment = 3

    def write_list(self, content_list):
        for content in content_list:
            self.document.add_paragraph(
            content, style='List Bullet'
        )
          
    def write_block(self,  title, content, bullet_list):
        paragraph = self.document.add_paragraph(
            ''
        ).add_run(title)
        paragraph.font.color.rgb = RGBColor.from_string('0000FF')

        # self.document.add_paragraph('').add_run(content).italic
        
        # self.write_paragraph(content)
        
        # self.write_list(bullet_list)

    def write_table(self, heading=[], records=[]):
        table = self.document.add_table(rows=1, cols=len(heading))
        table.style = 'Light Grid Accent 5'
        
        hdr_cells = table.rows[0].cells
        for column_idx in range(len(heading)):
            hdr_cells[column_idx].text = heading[column_idx]

        for row in records:
            row_cells = table.add_row().cells
            for column_idx in range(len(heading)):
                row_cells[column_idx].text = row[heading[column_idx]]
                
    def get_status_vietnamese(self, status):
        """Convert English status to Vietnamese"""
        status_map = {
            'new': 'Chưa tiếp nhận',
            'in_progress': 'Đang thực hiện',
            'pending_approval': 'Chờ phê duyệt',
            'completed': 'Hoàn thành',
            'overdue': 'Quá hạn',
            'paused': 'Tạm dừng',
            'cancelled': 'Đã hủy'
        }
        return status_map.get(status, status)
