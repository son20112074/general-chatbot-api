from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
import os
import re
from bs4 import BeautifulSoup
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.domain.models.task import Task
from app.domain.models.task_work import TaskWork
from app.domain.models.user import User


class TaskWorkDocxCreator:
    def __init__(self, session):
        self.session = session
    
    async def create_docx(self, work_list, download_options=[]):
        self.document = Document()
        self.write_title(
            "DANH SÁCH CÔNG VIỆC THỰC HIỆN", 
            "Thông tin chi tiết của {} công việc thực hiện".format(len(work_list))
        ) 
        self.write_paragraph("", '')
        
        # Create table for all works
        if work_list:
            # Prepare table data
            table_headers = ["STT", "Tiêu đề", "Mô tả", "Nội dung", "Người thực hiện", "Thời gian tạo", "Nhiệm vụ", "Trạng thái nhiệm vụ"]
            table_records = []
            
            for i, work in enumerate(work_list, 1):
                # Get work information
                work_title = work.get('title', 'Công việc không có tiêu đề')
                work_description = work.get('description', 'Không có mô tả')
                
                # Convert HTML content to text
                work_content = work.get('content', 'Không có nội dung')
                content_text = self.html_to_text(work_content) if work_content else "Không có nội dung"
                
                # Get creator information
                creator = work.get('creator')
                if creator:
                    creator_name = creator.get('full_name', 'Không xác định')
                    creator_account = creator.get('account_name', '')
                    creator_info = f"{creator_account} - {creator_name}"
                else:
                    creator_info = 'Không xác định'
                
                # Format creation time
                created_at = work.get('created_at')
                if created_at:
                    if isinstance(created_at, str):
                        try:
                            created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                        except:
                            pass
                    
                    if isinstance(created_at, datetime):
                        created_time_str = created_at.strftime('%d/%m/%Y %H:%M')
                    else:
                        created_time_str = str(created_at)
                else:
                    created_time_str = 'Không xác định'
                
                # Get task information
                task = work.get('task')
                if task:
                    task_id = task.get('id', 'Không có ID')
                    task_title = task.get('title', 'Không có tiêu đề')
                    task_info = f"(Mã {task_id}) - {task_title}"
                    
                    # Get task status
                    task_status = task.get('status', 'new')
                    status_vn = self.get_status_vietnamese(task_status)
                else:
                    task_info = 'Không có thông tin nhiệm vụ'
                    status_vn = 'Không xác định'
                
                # Add to table records
                table_records.append({
                    "STT": str(i),
                    "Tiêu đề": work_title,
                    "Mô tả": work_description,
                    "Nội dung": content_text,
                    "Người thực hiện": creator_info,
                    "Thời gian tạo": created_time_str,
                    "Nhiệm vụ": task_info,
                    "Trạng thái nhiệm vụ": status_vn
                })
            
            # Write the table
            self.write_table(table_headers, table_records)
        else:
            self.write_paragraph("Không có công việc nào để hiển thị", "")
        
        file_name = '{}_{}.docx'.format('danh_sach_cong_viec', str(datetime.today().timestamp()))
        
        # Create save path with year/month/day structure
        current_date = datetime.now()
        year = str(current_date.year)
        month = str(current_date.month).zfill(2)  # Zero-padded month
        day = str(current_date.day).zfill(2)      # Zero-padded day
        
        save_path = f'static/downloads/docx/{year}/{month}/{day}'
        
        # Create directory structure if it doesn't exist
        os.makedirs(save_path, exist_ok=True)

        self.document.save(save_path + '/' + file_name)
        
        return settings.BASE_URL + '/' + save_path + '/' + file_name



    def html_to_text(self, html_content):
        """Convert HTML content to plain text"""
        if not html_content:
            return ""
            
        try:
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            return text
            
        except Exception as e:
            print(f"Error converting HTML to text: {e}")
            # Fallback: return original content with basic HTML tag removal
            return re.sub(r'<[^>]+>', '', html_content)

    def set_empty_to_none(self, info={}):
        for key, value in info.items():
            if value == None:
                info[key] = ''
        # pass


    def write_title(self, title, summary):
        self.document.add_heading('', 0).add_run(title).font.size = Pt(16)
        p = self.document.add_paragraph('').add_run(summary).italic = True


    def write_heading_bullet(self, title):
        self.document.add_paragraph(
            title, style='List Number'
        )
        
    def write_paragraph(self, title_bold, content):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(title_bold).bold = True
        paragraph.add_run(content).italic = True
        # paragraph.alignment = 3

    def write_list(self, content_list):
        for content in content_list:
            self.document.add_paragraph(
            content, style='List Bullet'
        )
          
    def write_block(self,  title, content, bullet_list):
        paragraph = self.document.add_paragraph(
            ''
        ).add_run(title)
        paragraph.font.color.rgb = RGBColor.from_string('0000FF')

        # self.document.add_paragraph('').add_run(content).italic
        
        # self.write_paragraph(content)
        
        # self.write_list(bullet_list)

    def write_table(self, heading=[], records=[]):
        table = self.document.add_table(rows=1, cols=len(heading))
        table.style = 'Light Grid Accent 5'
        
        hdr_cells = table.rows[0].cells
        for column_idx in range(len(heading)):
            hdr_cells[column_idx].text = heading[column_idx]

        for row in records:
            row_cells = table.add_row().cells
            for column_idx in range(len(heading)):
                row_cells[column_idx].text = row[heading[column_idx]]
                
    def get_status_vietnamese(self, status):
        """Convert English status to Vietnamese"""
        status_map = {
            'new': 'Chưa tiếp nhận',
            'in_progress': 'Đang thực hiện',
            'pending_approval': 'Chờ phê duyệt',
            'completed': 'Hoàn thành',
            'overdue': 'Quá hạn',
            'paused': 'Tạm dừng',
            'cancelled': 'Đã hủy'
        }
        return status_map.get(status, status)


class TaskWorkDocxCreatorGroup:
    def __init__(self, session):
        self.session = session
    
    async def create_docx(self, work_list, download_options=[]):
        self.document = Document()
        self.write_title(
            "DANH SÁCH CÔNG VIỆC THỰC HIỆN (THEO NGƯỜI THỰC HIỆN)", 
            "Thông tin chi tiết của {} công việc thực hiện được nhóm theo người thực hiện".format(len(work_list))
        ) 
        self.write_paragraph("", '')
        
        # Group works by creator
        grouped_works = self.group_works_by_creator(work_list)
        
        # Process each creator group
        creator_index = 1
        for creator_key, works in grouped_works.items():
            await self.process_creator_group(creator_key, works, str(creator_index))
            creator_index += 1
        
        file_name = '{}_{}.docx'.format('danh_sach_cong_viec_theo_nguoi_thuc_hien', str(datetime.today().timestamp()))
        
        # Create save path with year/month/day structure
        current_date = datetime.now()
        year = str(current_date.year)
        month = str(current_date.month).zfill(2)  # Zero-padded month
        day = str(current_date.day).zfill(2)      # Zero-padded day
        
        save_path = f'static/downloads/docx/{year}/{month}/{day}'
        
        # Create directory structure if it doesn't exist
        os.makedirs(save_path, exist_ok=True)

        self.document.save(save_path + '/' + file_name)
        
        return settings.BASE_URL + '/' + save_path + '/' + file_name

    def group_works_by_creator(self, work_list):
        """Group works by creator"""
        grouped = {}
        
        for work in work_list:
            creator = work.get('creator')
            if creator:
                creator_name = creator.get('full_name', 'Không xác định')
                creator_account = creator.get('account_name', '')
                creator_key = f"{creator_account} - {creator_name}"
            else:
                creator_key = "Không xác định"
            
            if creator_key not in grouped:
                grouped[creator_key] = []
            
            grouped[creator_key].append(work)
        
        return grouped

    async def process_creator_group(self, creator_key, works, creator_index):
        """Process a group of works by the same creator"""
        print(f"Processing creator group {creator_index}: {creator_key} with {len(works)} works")
        
        # Write creator header
        creator_title = f"{creator_index}. {creator_key}"
        self.write_block(creator_title, '', [])
        
        # Write creator information (common for all works in this group)
        if works and works[0].get('creator'):
            creator = works[0].get('creator')
            creator_name = creator.get('full_name', 'Không xác định')
            creator_account = creator.get('account_name', '')
            creator_email = creator.get('email', '')
            
            self.write_paragraph('Tên đăng nhập: ', creator_account)
            self.write_paragraph('Họ và tên: ', creator_name)
            if creator_email:
                self.write_paragraph('Email: ', creator_email)
            
            # Add separator after creator info
            self.write_paragraph("", "")
        
        # Create table for works in this group
        if works:
            # Prepare table data
            table_headers = ["STT", "Tiêu đề", "Mô tả", "Nội dung", "Thời gian tạo", "Nhiệm vụ", "Trạng thái nhiệm vụ"]
            table_records = []
            
            for i, work in enumerate(works, 1):
                work_index = f"{creator_index}.{i}"
                
                # Get work information
                work_title = work.get('title', 'Công việc không có tiêu đề')
                work_description = work.get('description', 'Không có mô tả')
                
                # Convert HTML content to text
                work_content = work.get('content', 'Không có nội dung')
                content_text = self.html_to_text(work_content) if work_content else "Không có nội dung"
                
                # Format creation time
                created_at = work.get('created_at')
                if created_at:
                    if isinstance(created_at, str):
                        try:
                            created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                        except:
                            pass
                    
                    if isinstance(created_at, datetime):
                        created_time_str = created_at.strftime('%d/%m/%Y %H:%M')
                    else:
                        created_time_str = str(created_at)
                else:
                    created_time_str = 'Không xác định'
                
                # Get task information
                task = work.get('task')
                if task:
                    task_id = task.get('id', 'Không có ID')
                    task_title = task.get('title', 'Không có tiêu đề')
                    task_info = f"(Mã {task_id}) - {task_title}"
                    
                    # Get task status
                    task_status = task.get('status', 'new')
                    status_vn = self.get_status_vietnamese(task_status)
                else:
                    task_info = 'Không có thông tin nhiệm vụ'
                    status_vn = 'Không xác định'
                
                # Add to table records
                table_records.append({
                    "STT": work_index,
                    "Tiêu đề": work_title,
                    "Mô tả": work_description,
                    "Nội dung": content_text,
                    "Thời gian tạo": created_time_str,
                    "Nhiệm vụ": task_info,
                    "Trạng thái nhiệm vụ": status_vn
                })
            
            # Write the table
            self.write_table(table_headers, table_records)
        else:
            self.write_paragraph("Không có công việc nào để hiển thị", "")
        
        # Add separator between creator groups
        self.write_paragraph("", "")



    def html_to_text(self, html_content):
        """Convert HTML content to plain text"""
        if not html_content:
            return ""
            
        try:
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            return text
            
        except Exception as e:
            print(f"Error converting HTML to text: {e}")
            # Fallback: return original content with basic HTML tag removal
            return re.sub(r'<[^>]+>', '', html_content)

    def write_title(self, title, summary):
        self.document.add_heading('', 0).add_run(title).font.size = Pt(16)
        p = self.document.add_paragraph('').add_run(summary).italic = True

    def write_heading_bullet(self, title):
        self.document.add_paragraph(
            title, style='List Number'
        )
        
    def write_paragraph(self, title_bold, content):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(title_bold).bold = True
        paragraph.add_run(content).italic = True

    def write_list(self, content_list):
        for content in content_list:
            self.document.add_paragraph(
            content, style='List Bullet'
        )
          
    def write_block(self, title, content, bullet_list):
        paragraph = self.document.add_paragraph(
            ''
        ).add_run(title)
        paragraph.font.color.rgb = RGBColor.from_string('0000FF')

    def write_table(self, heading=[], records=[]):
        table = self.document.add_table(rows=1, cols=len(heading))
        table.style = 'Light Grid Accent 5'
        
        hdr_cells = table.rows[0].cells
        for column_idx in range(len(heading)):
            hdr_cells[column_idx].text = heading[column_idx]

        for row in records:
            row_cells = table.add_row().cells
            for column_idx in range(len(heading)):
                row_cells[column_idx].text = row[heading[column_idx]]
                
    def get_status_vietnamese(self, status):
        """Convert English status to Vietnamese"""
        status_map = {
            'new': 'Chưa tiếp nhận',
            'in_progress': 'Đang thực hiện',
            'pending_approval': 'Chờ phê duyệt',
            'completed': 'Hoàn thành',
            'overdue': 'Quá hạn',
            'paused': 'Tạm dừng',
            'cancelled': 'Đã hủy'
        }
        return status_map.get(status, status)
