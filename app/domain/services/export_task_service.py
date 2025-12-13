from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
import os
import re
from bs4 import BeautifulSoup
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.domain.models.task import Task
from app.domain.models.task_work import TaskWork
from app.domain.models.user import User


class TaskDocxCreator:
    def __init__(self, session):
        self.session = session
    
    async def create_docx(self, task_list, download_options=[]):
        self.document = Document()
        self.write_title(
            "DANH SÁCH NHIỆM VỤ", 
            "Thông tin chi tiết của {} nhiệm vụ và các công việc thực hiện".format(len(task_list))
        ) 
        self.write_paragraph("", '')
        
        # Process each task in the tree structure with hierarchical numbering
        for i, task in enumerate(task_list, 1):
            await self.process_task(task, str(i))
        
        file_name = '{}_{}.docx'.format('danh_sach_nhiem_vu', str(datetime.today().timestamp()))
        
        # Create save path with year/month/day structure
        current_date = datetime.now()
        year = str(current_date.year)
        month = str(current_date.month).zfill(2)  # Zero-padded month
        day = str(current_date.day).zfill(2)      # Zero-padded day
        
        save_path = f'static/downloads/docx/{year}/{month}/{day}'
        
        # Create directory structure if it doesn't exist
        os.makedirs(save_path, exist_ok=True)

        self.document.save(save_path + '/' + file_name)
        
        return settings.BASE_URL + '/' + save_path + '/' + file_name

    async def process_task(self, task, index=""):
        """Process a single task and its children recursively with hierarchical numbering"""
                
        # Write task title and description with hierarchical index
        task_title = f"{index}. {task.get('title', 'Nhiệm vụ không có tiêu đề')}"
        task_description = task.get('description', 'Không có mô tả')
        
        self.write_block(task_title, '',[])
        # Use write_block for the task title and description
        self.write_paragraph('Mô tả: ', task_description)
        
        # Add start and end dates
        start_date = task.get('start_date')
        end_date = task.get('end_date')
        
        if start_date:
            # Convert ISO string to datetime if needed
            if isinstance(start_date, str):
                try:
                    start_date = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
                except:
                    pass
            
            if isinstance(start_date, datetime):
                start_date_str = start_date.strftime('%d/%m/%Y %H:%M')
            else:
                start_date_str = str(start_date)
            self.write_paragraph('Thời gian bắt đầu: ', start_date_str)
        else:
            self.write_paragraph('Thời gian bắt đầu: ', 'Chưa thiết lập')
            
        if end_date:
            # Convert ISO string to datetime if needed
            if isinstance(end_date, str):
                try:
                    end_date = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
                except:
                    pass
                    
            if isinstance(end_date, datetime):
                end_date_str = end_date.strftime('%d/%m/%Y %H:%M')
            else:
                end_date_str = str(end_date)
            self.write_paragraph('Thời gian kết thúc: ', end_date_str)
        else:
            self.write_paragraph('Thời gian kết thúc: ', 'Chưa thiết lập')
        
        # Add status with Vietnamese translation
        status = task.get('status', 'new')
        status_vn = self.get_status_vietnamese(status)
        self.write_paragraph('Trạng thái: ', status_vn)
        
        # Query and write task works for this task
        task_id = task.get('id')
        await self.write_task_works(task_id)
        
        # Process children tasks recursively with hierarchical numbering
        children = task.get('children', [])
        for i, child in enumerate(children, 1):
            child_index = f"{index}.{i}" if index else str(i)
            await self.process_task(child, child_index)

    async def write_task_works(self, task_id):
        """Query and write task works for a specific task"""
        if not task_id:
            self.write_paragraph("Công việc thực hiện: ", "Không có ID nhiệm vụ")
            return
            
        try:
            # Convert task_id to integer if it's a string
            if isinstance(task_id, str):
                task_id = int(task_id)
                        
            # Query task works from database with user information
            query = select(TaskWork, User).join(
                User, TaskWork.created_by == User.id
            ).where(
                TaskWork.task_id == task_id
            ).order_by(TaskWork.created_at.desc())
            
            result = await self.session.execute(query)
            task_works_with_users = result.all()
                        
            if task_works_with_users:
                self.write_paragraph("Danh sách công việc thực hiện:", "")
                
                # Prepare table data
                table_headers = ["STT", "Tiêu đề", "Mô tả", "Nội dung", "Người thực hiện", "Thời gian tạo"]
                table_records = []
                
                for i, (task_work, user) in enumerate(task_works_with_users, 1):
                    work_title = task_work.title
                    work_description = task_work.description or "Không có mô tả"
                    
                    # Convert HTML content to text
                    work_content = self.html_to_text(task_work.content) if task_work.content else "Không có nội dung"
                    
                    # Get user information
                    user_name = user.full_name if user else "Không xác định"
                    
                    # Format creation time
                    created_time = task_work.created_at.strftime('%d/%m/%Y %H:%M') if task_work.created_at else "Không xác định"
                    
                    # Add to table records
                    table_records.append({
                        "STT": str(i),
                        "Tiêu đề": work_title,
                        "Mô tả": work_description,
                        "Nội dung": work_content,
                        "Người thực hiện": user_name,
                        "Thời gian tạo": created_time
                    })
                
                # Write the table
                self.write_table(table_headers, table_records)
                        
            else:
                self.write_paragraph("Công việc thực hiện: ", "Chưa có công việc nào được tạo")
                
        except ValueError as e:
            print(f"ValueError converting task_id {task_id}: {e}")
            self.write_paragraph("Công việc thực hiện: ", f"Lỗi ID nhiệm vụ: {task_id}")
        except Exception as e:
            print(f"Error querying task works for task {task_id}: {e}")
            print(f"Error type: {type(e)}")
            import traceback
            traceback.print_exc()
            self.write_paragraph("Công việc thực hiện: ", "Lỗi khi truy vấn dữ liệu")

    def html_to_text(self, html_content):
        """Convert HTML content to plain text"""
        if not html_content:
            return ""
            
        try:
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            return text
            
        except Exception as e:
            print(f"Error converting HTML to text: {e}")
            # Fallback: return original content with basic HTML tag removal
            return re.sub(r'<[^>]+>', '', html_content)

    def set_empty_to_none(self, info={}):
        for key, value in info.items():
            if value == None:
                info[key] = ''
        # pass


    def write_title(self, title, summary):
        self.document.add_heading('', 0).add_run(title).font.size = Pt(16)
        p = self.document.add_paragraph('').add_run(summary).italic = True


    def write_heading_bullet(self, title):
        self.document.add_paragraph(
            title, style='List Number'
        )
        
    def write_paragraph(self, title_bold, content):
        paragraph = self.document.add_paragraph()
        paragraph.add_run(title_bold).bold = True
        paragraph.add_run(content).italic = True
        # paragraph.alignment = 3

    def write_list(self, content_list):
        for content in content_list:
            self.document.add_paragraph(
            content, style='List Bullet'
        )
          
    def write_block(self,  title, content, bullet_list):
        paragraph = self.document.add_paragraph(
            ''
        ).add_run(title)
        paragraph.font.color.rgb = RGBColor.from_string('0000FF')

        # self.document.add_paragraph('').add_run(content).italic
        
        # self.write_paragraph(content)
        
        # self.write_list(bullet_list)

    def write_table(self, heading=[], records=[]):
        table = self.document.add_table(rows=1, cols=len(heading))
        table.style = 'Light Grid Accent 5'
        
        hdr_cells = table.rows[0].cells
        for column_idx in range(len(heading)):
            hdr_cells[column_idx].text = heading[column_idx]

        for row in records:
            row_cells = table.add_row().cells
            for column_idx in range(len(heading)):
                row_cells[column_idx].text = row[heading[column_idx]]
                
    def get_status_vietnamese(self, status):
        """Convert English status to Vietnamese"""
        status_map = {
            'new': 'Chưa tiếp nhận',
            'in_progress': 'Đang thực hiện',
            'pending_approval': 'Chờ phê duyệt',
            'completed': 'Hoàn thành',
            'overdue': 'Quá hạn',
            'paused': 'Tạm dừng',
            'cancelled': 'Đã hủy'
        }
        return status_map.get(status, status)
